# utils.py
"""
LexPrep utilities
-----------------
Currently provides one helper:

    extract_placeholders(docx_path)  →  list[str]

It scans a Word document for double-brace placeholders such as:

    {{ plaintiff_name }}
    {{ signature_date }}
    {{ plaintiffs[].address }}

and returns a **sorted, de-duplicated** list of the raw keys
(e.g., "plaintiff_name", "plaintiffs[].address").

Placeholders with the pattern  `root[].subkey`  are intended to signal
“repeat” groups (tables or lists) when the manifest is auto-built in
app.py, but this module itself just reports the strings it finds.

Requires `python-docx`:
    pip install python-docx
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Set


import os
import json
import openai
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()  # so OPENAI_API_KEY is picked up from your .env
openai.api_key = os.getenv("OPENAI_API_KEY")


def extract_fields_from_text(text: str) -> list[dict]:
    """
    Send document text to OpenAI and parse out fillable fields.
    Returns a list of {"name": "...", "type": "text"/"date"/"number"}.
    """
    prompt = (
        "You are a legal-tech assistant. Given the text of a legal document, "
        "identify all the distinct fillable fields (e.g., party names, dates, amounts). "
        "Return a JSON array of objects with keys: "
        "  - name (camelCase identifier), "
        "  - type (text, date, number).\n\n"
        "Document excerpt:\n"
        f"{text[:3000]}…"
    )

    resp = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Extract template fields from legal text."},
            {"role": "user",   "content": prompt}
        ],
        temperature=0,
    )
    return json.loads(resp.choices[0].message.content)


def save_template_file(filename: str, data: bytes) -> str:
    """
    Save the raw upload under default_templates/ and return its path.
    """
    tmpl_dir = Path(__file__).parent / "default_templates"
    tmpl_dir.mkdir(parents=True, exist_ok=True)
    dst = tmpl_dir / filename
    with open(dst, "wb") as f:
        f.write(data)
    return str(dst)




from docx import Document  # python-docx


# Matches {{ some_key }}  where the key may contain letters, numbers,
# underscores, dots, and square brackets (for repeat groups).
_FIELD_RE = re.compile(r"{{\s*([a-zA-Z0-9_.\[\]]+)\s*}}")


def extract_placeholders(docx_path: Path | str) -> List[str]:
    """
    Scan the DOCX at *docx_path* and return a sorted list of unique
    placeholder keys found anywhere in the document (paragraphs and tables).

    Parameters
    ----------
    docx_path : pathlib.Path | str
        Path to the .docx file to inspect.

    Returns
    -------
    List[str]
        Sorted list of unique placeholder keys.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    document = Document(docx_path)
    keys: Set[str] = set()

    def _collect(text: str) -> None:
        if not text:
            return
        keys.update(_FIELD_RE.findall(text))

    # scan paragraphs
    for para in document.paragraphs:
        _collect(para.text)

    # scan every cell in every table
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _collect(cell.text)

    return sorted(keys)
