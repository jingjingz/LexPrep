# utils.py

import os
import re
import json
import openai
from pathlib import Path
from dotenv import load_dotenv
from typing import List
import docx  # for your original placeholder extractor

# ── Load environment & configure OpenAI ───────────────────────────────────────
load_dotenv()  # reads .env at project root
openai.api_key = os.getenv("OPENAI_API_KEY")

import tempfile
from docx import Document
import pypandoc

def extract_text_from_bytes(raw_bytes: bytes, file_mime: str) -> str:
    """
    Write the uploaded bytes to a temp file (either .docx or .pdf),
    then extract and return its plain-text contents.
    """
    # choose extension from mime
    ext = "pdf" if "pdf" in file_mime else "docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(raw_bytes)
        tmp_path = tmp.name

    if ext == "docx":
        # use python-docx
        doc = Document(tmp_path)
        lines = [para.text for para in doc.paragraphs]
        text = "\n".join(lines)
    else:
        # use pypandoc to convert PDF → plain
        text = pypandoc.convert_file(tmp_path, "plain")

    return text




# ── 1) Your original DOCX‐placeholder extractor ──────────────────────────────
_FIELD_RE = re.compile(r"\{\{\s*([^{}\s]+)\s*\}\}")

def extract_placeholders(docx_path: str) -> List[str]:
    """
    Scan a DOCX for double-brace placeholders and return
    a sorted, de-duplicated list of the raw keys.
    """
    def _collect(text: str):
        if text:
            keys.update(_FIELD_RE.findall(text))

    document = docx.Document(docx_path)
    keys = set()

    for para in document.paragraphs:
        _collect(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _collect(cell.text)

    return sorted(keys)


# ── 2) OpenAI‐powered field extractor ─────────────────────────────────────────
def extract_fields_from_text(text: str) -> List[dict]:
    """
    Send document text to OpenAI and parse out fillable fields.
    Returns a list of {"name": "...", "type": "text"/"date"/"number"}.
    """
    prompt = (
        "You are a legal-tech assistant. Given the text of a legal document, "
        "identify all distinct fillable fields (e.g., party names, dates, amounts). "
        "Return a JSON array of objects with keys:\n"
        "  - name: a camelCase identifier\n"
        "  - type: one of text, date, or number\n\n"
        "Document excerpt:\n"
        f"{text[:3000]}…"
    )

    resp = openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Extract template fields from legal text."},
            {"role": "user",   "content": prompt}
        ],
        temperature=0,
    )

    # The shape of the response is the same, so parsing is unchanged
    return json.loads(resp.choices[0].message.content)



# ── 3) Save raw uploads for templates ──────────────────────────────────────────
def save_template_file(filename: str, data: bytes) -> str:
    """
    Save the raw upload under default_templates/ and return its path.
    """
    tmpl_dir = Path(__file__).parent / "default_templates"
    tmpl_dir.mkdir(parents=True, exist_ok=True)

    dst = tmpl_dir / filename
    with open(dst, "wb") as f:
        f.write(data)

    return str(dst)
